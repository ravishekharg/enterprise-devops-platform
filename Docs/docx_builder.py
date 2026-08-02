"""Shared helper for building HLD/LLD .docx files in the same style as the
existing Darviq-Health reference docs (Docs/Virenza_Health_High_Level_Design.docx
and _Low_Level_Design.docx). Copy this file next to your generation script and
import it — don't reinvent the docx structure per repo, so all 20 doc pairs
stay visually/structurally consistent.

Usage sketch (see bottom of file for a runnable example):

    from docx_builder import DesignDoc

    doc = DesignDoc(
        project_name="Darviq Fleet",
        subtitle="GitOps Multi-Cluster Kubernetes Fleet Management",
        doc_kind="High-Level Design (HLD)",
        version="1.0",
        date="July 31, 2026",
    )
    doc.add_document_control()
    doc.add_toc_field()

    doc.add_heading1("1. Introduction")
    doc.add_heading2("1.1 Purpose")
    doc.add_paragraph("This document describes ...")
    doc.add_heading2("1.2 Scope")
    doc.add_bullets(["Thing one", "Thing two"])

    doc.add_heading1("3. Architecture overview")
    doc.add_table(
        headers=["Component", "Responsibility", "Technology"],
        rows=[["ArgoCD", "GitOps reconciliation", "ArgoCD"], ...],
    )

    doc.save("Docs/Darviq-Fleet_High_Level_Design.docx")
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Resolved relative to this file (not the caller's cwd) so every project's
# own copy of docx_builder.py finds the darviq-logo.png sitting next to it
# in that project's Docs/ folder, regardless of where the generator script
# is invoked from.
DEFAULT_LOGO_PATH = str(Path(__file__).resolve().parent / "darviq-logo.png")


class DesignDoc:
    def __init__(
        self,
        project_name: str,
        subtitle: str,
        doc_kind: str,
        version: str,
        date: str,
        status: str = "Version 1.0",
        logo_path: str | None = DEFAULT_LOGO_PATH,
    ):
        self.doc = Document()
        self._style_base()
        self._title_page(project_name, subtitle, doc_kind, version, date, status, logo_path)
        self.project_name = project_name
        self.doc_kind = doc_kind
        self.version = version
        self.date = date
        self.status = status

    # ---- setup -----------------------------------------------------
    def _style_base(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

    def add_logo(self, logo_path: str | None, width_inches: float = 1.3) -> None:
        if not logo_path or not os.path.exists(logo_path):
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(width_inches))

    def _title_page(self, project_name, subtitle, doc_kind, version, date, status, logo_path=DEFAULT_LOGO_PATH) -> None:
        d = self.doc
        self.add_logo(logo_path)

        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(project_name)
        run.font.size = Pt(28)
        run.font.bold = True

        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{doc_kind} Document")
        run.font.size = Pt(18)
        run.font.bold = True

        for text in (f"Version {version}", date, f"Status: {status}"):
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(text)

        d.add_page_break()

    # ---- structural building blocks ---------------------------------
    def add_document_control(self, extra_rows: list[tuple[str, str]] | None = None) -> None:
        self.add_heading1("Document control")
        rows = [
            ("Document title", f"{self.project_name} — {self.doc_kind}"),
            ("Version", self.version),
            ("Date", self.date),
            ("Status", self.status),
            ("Author", "Darviq Systems engineering"),
            ("Owner", "Darviq Systems engineering"),
        ]
        if extra_rows:
            rows.extend(extra_rows)
        self.add_table(headers=["Field", "Value"], rows=[list(r) for r in rows])

    def add_toc_field(self) -> None:
        """Inserts a real Word TOC field (right-click -> Update Field after
        opening in Word to populate it, or Word may prompt automatically)."""
        self.add_heading1("Table of contents")
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(fld_char_end)
        self.doc.add_page_break()

    def add_heading1(self, text: str):
        return self.doc.add_heading(text, level=1)

    def add_heading2(self, text: str):
        return self.doc.add_heading(text, level=2)

    def add_heading3(self, text: str):
        return self.doc.add_heading(text, level=3)

    def add_paragraph(self, text: str):
        return self.doc.add_paragraph(text)

    def add_bullets(self, items: list[str]) -> None:
        for item in items:
            self.doc.add_paragraph(item, style="List Bullet")

    def add_code_block(self, text: str) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)

    def add_table(self, headers: list[str], rows: list[list[str]]) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        self.doc.add_paragraph()  # spacing after table

    def add_figure_placeholder(self, caption: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[{caption}]")
        run.italic = True

    def save(self, path: str) -> None:
        self.doc.save(path)


if __name__ == "__main__":
    # Smoke test
    doc = DesignDoc(
        project_name="Example Project",
        subtitle="Example Subtitle",
        doc_kind="High-Level Design (HLD)",
        version="1.0",
        date="July 31, 2026",
    )
    doc.add_document_control()
    doc.add_toc_field()
    doc.add_heading1("1. Introduction")
    doc.add_heading2("1.1 Purpose")
    doc.add_paragraph("Example purpose text.")
    doc.add_table(headers=["A", "B"], rows=[["1", "2"], ["3", "4"]])
    doc.save("smoke_test.docx")
    print("OK")
